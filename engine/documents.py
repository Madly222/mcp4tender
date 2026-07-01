from __future__ import annotations

import base64
import hashlib
import urllib.request
from pathlib import Path

USER_AGENT = "tenderengine/1.0 (+https://rapidlink.md)"


def _cache_path(cache_dir, url):
    h = hashlib.sha256(url.encode("utf-8")).hexdigest()[:32]
    return Path(cache_dir) / h


def fetch_document(url, cache_dir, timeout=60):
    path = _cache_path(cache_dir, url)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists() and path.stat().st_size > 0:
        return str(path)
    req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        data = resp.read()
    path.write_bytes(data)
    return str(path)


def extract_pdf_text(path):
    import fitz
    doc = fitz.open(path)
    texts = [page.get_text() for page in doc]
    pages = doc.page_count
    doc.close()
    return "\n".join(texts), pages


def extract_docx_text(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts), max(len(d.paragraphs), 1)


def is_scanned(text, pages, threshold=100):
    pages = max(pages, 1)
    return (len(text.strip()) / pages) < threshold


def render_pages_png(path, max_pages=5, zoom=2.0):
    import fitz
    doc = fitz.open(path)
    images = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        images.append(pix.tobytes("png"))
    doc.close()
    return images


def ocr_via_vision(images, gateway, config):
    blocks = []
    for img in images:
        b64 = base64.standard_b64encode(img).decode("ascii")
        blocks.append({"type": "image",
                       "source": {"type": "base64", "media_type": "image/png", "data": b64}})
    prompt = config.get("documents.ocr_prompt",
                        "Transcrie integral textul din acest document. Raspunde doar cu textul.")
    blocks.append({"type": "text", "text": prompt})
    r = gateway.complete("ocr", "", [{"role": "user", "content": blocks}],
                         max_tokens=int(config.get("documents.ocr_max_tokens", 4096)))
    return r["text"], r


def get_document_text(doc, gateway, config):
    fmt = (doc.get("format") or "").lower()
    url = doc.get("url")
    cache_dir = config.get("documents.cache_dir", "data/docs")
    skip = config.get("documents.skip_formats", ["doc"])

    if not url:
        return {"text": "", "method": "skipped", "error": "no url"}
    if fmt in skip:
        return {"text": "", "method": "skipped", "error": f"format {fmt}"}

    try:
        path = fetch_document(url, cache_dir, int(config.get("documents.timeout", 60)))
    except Exception as exc:
        return {"text": "", "method": "fetch_error", "error": str(exc)}

    try:
        if fmt == "pdf":
            text, pages = extract_pdf_text(path)
        elif fmt == "docx":
            text, pages = extract_docx_text(path)
        else:
            return {"text": "", "method": "unsupported", "error": f"format {fmt}"}
    except Exception as exc:
        return {"text": "", "method": "extract_error", "error": str(exc)}

    method = "text"
    cost = 0.0
    tokens = 0

    if fmt == "pdf" and config.get("documents.vision_enabled", True) and \
            is_scanned(text, pages, int(config.get("documents.scanned_threshold", 100))):
        try:
            images = render_pages_png(path, int(config.get("documents.max_pages_ocr", 5)))
            vtext, r = ocr_via_vision(images, gateway, config)
            if len(vtext.strip()) > len(text.strip()):
                text, method = vtext, "vision"
                cost = r.get("cost", 0)
                tokens = r.get("input_tokens", 0) + r.get("output_tokens", 0)
        except Exception as exc:
            method = "text_ocr_failed"

    return {"text": text, "method": method, "pages": pages,
            "chars": len(text), "cost": cost, "tokens": tokens}
